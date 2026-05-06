# -*- coding: utf-8 -*-
from odoo import models, fields, api, _
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import base64
import io
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

class HopDongLaoDong(models.Model):
    _name = 'nhan_su.hop_dong'
    _inherit = ['mail.thread', 'mail.activity.mixin']
    _description = 'Hợp đồng lao động'
    _order = 'ngay_ky desc'

    name = fields.Char(string='Số hợp đồng', required=True, copy=False, default='/')
    nhan_vien_id = fields.Many2one('nhan_su.nhan_vien', string='Nhân viên', ondelete='cascade')
    ngay_ky = fields.Date(string='Ngày ký', default=fields.Date.context_today)
    ngay_bat_dau = fields.Date(string='Ngày bắt đầu')
    ngay_ket_thuc = fields.Date(string='Ngày kết thúc')
    company_id = fields.Many2one('res.company', string='Công ty', default=lambda self: self.env.company)
    chuc_vu_id = fields.Many2one('nhan_su.chuc_vu', string='Chức vụ', related='nhan_vien_id.chuc_vu_id', store=True, readonly=True)
    state = fields.Selection([
        ('draft', 'Mới tạo'),
        ('active', 'Đang hiệu lực'),
        ('expiring', 'Sắp hết hạn'),
        ('expired', 'Hết hạn'),
        ('cancel', 'Đã hủy'),
    ], string='Trạng thái', default='draft', tracking=True)

    loai_hop_dong = fields.Selection([
        ('36_thang', 'Có thời hạn 36 tháng'),
        ('thu_viec', 'Thử việc (85%)'),
        ('chinh_thuc', 'Chính thức (100%)'),
        ('khac', 'Khác'),
    ], string='Loại hợp đồng', default='36_thang')

    # Thông tin Bên A (Công ty TNHH Thiên Thời)
    dia_chi_a = fields.Char(string='Địa chỉ Bên A', default='290 Huỳnh Thị Hai')
    msdn_a = fields.Char(string='Mã số doanh nghiệp', default='0110329220')
    nguoi_dai_dien_a = fields.Char(string='Người đại diện Bên A', default='Trần Đại Phú')
    chuc_vu_a = fields.Char(string='Chức vụ người đại diện', default='Giám đốc')
    phone_a = fields.Char(string='Số điện thoại Bên A')
    fax_a = fields.Char(string='Fax')
    mst_a = fields.Char(string='Mã số thuế Bên A', default='0110329220')

    # Địa điểm và ngày ký
    dia_diem_ky = fields.Char(string='Địa điểm ký', default='Thành phố Hồ Chí Minh')
    ngay_ky_text = fields.Char(compute='_compute_ngay_ky_text')

    @api.depends('ngay_ky', 'dia_diem_ky')
    def _compute_ngay_ky_text(self):
        for rec in self:
            if rec.ngay_ky:
                rec.ngay_ky_text = f'{rec.dia_diem_ky or "Thành phố Hồ Chí Minh"}, ngày {rec.ngay_ky.day:02d} tháng {rec.ngay_ky.month:02d} năm {rec.ngay_ky.year}'
            else:
                rec.ngay_ky_text = f'{rec.dia_diem_ky or "Thành phố Hồ Chí Minh"}, ngày ... tháng ... năm 20...'

    # Chi tiết công việc (Điều 1)
    dia_diem_lam_viec = fields.Char(string='Địa điểm làm việc', default='Trụ sở công ty')
    cong_viec_phai_lam = fields.Text(string='Công việc phải làm', default='Làm bánh, nhặt bánh, trộn bánh, ...')
    
    # Thông tin liên quan từ nhân viên (Cho phép chỉnh sửa trực tiếp)
    ngay_sinh_nv = fields.Date(related='nhan_vien_id.ngay_sinh', string='Ngày sinh NV', readonly=False)
    gioi_tinh_nv = fields.Selection(related='nhan_vien_id.gioi_tinh', string='Giới tính NV', readonly=False)
    so_cmnd_nv = fields.Char(related='nhan_vien_id.so_cmnd', string='Số CMND/CCCD NV', readonly=False)
    ngay_cap_cmnd_nv = fields.Date(related='nhan_vien_id.ngay_cap_cmnd', string='Ngày cấp NV', readonly=False)
    noi_cap_cmnd_nv = fields.Char(related='nhan_vien_id.noi_cap_cmnd', string='Nơi cấp NV', readonly=False)
    ho_khau_thuong_tru_nv = fields.Char(related='nhan_vien_id.ho_khau_thuong_tru', string='HKTT NV', readonly=False)
    dia_chi_lien_he_nv = fields.Char(related='nhan_vien_id.dia_chi_lien_he', string='Địa chỉ hiện nay NV', readonly=False)

    # Thời giờ làm việc (Điều 3)
    thoi_gio_lam_viec = fields.Char(string='Thời giờ làm việc', compute='_compute_working_hours', readonly=True)
    thoi_gian_nghi = fields.Char(string='Thời gian nghỉ', default='Hàng tuần: được nghỉ ngày chủ nhật')

    @api.depends('nhan_vien_id.ca_lam_id')
    def _compute_working_hours(self):
        for rec in self:
            ca = rec.nhan_vien_id.ca_lam_id if rec.nhan_vien_id else False
            if ca:
                h_sang_tu = f"{int(ca.sang_tu)}h{int((ca.sang_tu%1)*60):02d}"
                h_sang_den = f"{int(ca.sang_den)}h{int((ca.sang_den%1)*60):02d}"
                h_chieu_tu = f"{int(ca.chieu_tu)}h{int((ca.chieu_tu%1)*60):02d}"
                h_chieu_den = f"{int(ca.chieu_den)}h{int((ca.chieu_den%1)*60):02d}"
                rec.thoi_gio_lam_viec = f"8 tiếng/ngày – Ca sáng từ {h_sang_tu} đến {h_sang_den}, Ca chiều từ {h_chieu_tu} đến {h_chieu_den}"
            else:
                rec.thoi_gio_lam_viec = "8 tiếng/ngày – Ca sáng từ 7h00 đến 15h00, Ca chiều từ 15h00 đến 23h00"

    thoi_gian_nghi = fields.Char(string='Thời gian nghỉ', default='Hàng tuần: được nghỉ ngày chủ nhật')

    # ---------------------------------------------------------------
    # Lương cơ bản — tự động lấy từ Cấu hình tính lương
    # ---------------------------------------------------------------
    luong_gio_cau_hinh = fields.Float(
        string='Lương giờ (từ cấu hình)',
        compute='_compute_luong_tu_cau_hinh',
        store=True,
        help='Mức lương giờ lấy từ Cấu hình tính lương theo chức vụ nhân viên'
    )
    he_so_loai_hd = fields.Float(
        string='Hệ số loại HĐ',
        compute='_compute_luong_tu_cau_hinh',
        store=True,
        help='1.0 = Chính thức (100%), 0.85 = Thử việc (85%)'
    )
    luong_co_ban = fields.Float(
        string='Lương cơ bản (VNĐ/tháng)',
        compute='_compute_luong_tu_cau_hinh',
        store=True,
        readonly=False,
        help='Tự động tính = Lương giờ × Giờ/ngày × Ngày/tháng × Hệ số loại HĐ'
    )
    cong_thuc_tinh_luong = fields.Char(
        string='Công thức',
        compute='_compute_luong_tu_cau_hinh',
        store=True,
        help='Hiển thị công thức tính lương cơ bản'
    )

    @api.depends(
        'nhan_vien_id', 'nhan_vien_id.chuc_vu_id', 'nhan_vien_id.chuc_vu_id.name',
        'loai_hop_dong'
    )
    def _compute_luong_tu_cau_hinh(self):
        """Tự động tính lương cơ bản từ Cấu hình tính lương"""
        # Dùng .get() để tránh KeyError khi tinh_luong chưa được cài
        CauHinh = self.env.get('tinh_luong.cau_hinh')
        config = CauHinh.search([], limit=1) if CauHinh is not None else None

        for rec in self:
            if not config:
                rec.luong_gio_cau_hinh = 0.0
                rec.he_so_loai_hd = 1.0
                rec.luong_co_ban = 0.0
                rec.cong_thuc_tinh_luong = 'Chưa có cấu hình lương'
                continue

            # 1. Xác định lương giờ theo chức vụ
            chuc_vu = rec.nhan_vien_id.chuc_vu_id if rec.nhan_vien_id else False
            ten_chuc_vu = (chuc_vu.name or '').lower() if chuc_vu else ''

            # Nhận diện: nếu tên chức vụ chứa 'quản lý' / 'manager' → dùng lương QL
            if any(kw in ten_chuc_vu for kw in ['quản lý', 'quan ly', 'manager', 'giám đốc', 'truong', 'trưởng']):
                luong_gio = config.muc_luong_quan_ly
                loai_cv = 'Quản lý'
            else:
                luong_gio = config.muc_luong_nhan_vien
                loai_cv = 'Nhân viên'

            # 2. Hệ số theo loại hợp đồng
            if rec.loai_hop_dong == 'thu_viec':
                he_so = 0.85
                ten_loai = 'Thử việc (85%)'
            elif rec.loai_hop_dong == 'chinh_thuc':
                he_so = 1.0
                ten_loai = 'Chính thức (100%)'
            else:
                he_so = 1.0
                ten_loai = 'Khác (100%)'

            # 3. Tính lương cơ bản/tháng
            so_gio_ngay = config.so_gio_chuan_ngay or 8.0
            so_ngay_thang = config.so_ngay_cong_chuan_thang or 26.0
            luong_thang = luong_gio * so_gio_ngay * so_ngay_thang * he_so

            rec.luong_gio_cau_hinh = luong_gio
            rec.he_so_loai_hd = he_so
            rec.luong_co_ban = luong_thang
            rec.cong_thuc_tinh_luong = (
                f'{loai_cv}: {luong_gio:,.0f} VNĐ/giờ × {so_gio_ngay:.0f}h × {so_ngay_thang:.0f} ngày × {he_so:.0%} ({ten_loai}) = {luong_thang:,.0f} VNĐ'
            )

    # ---------------------------------------------------------------
    # Phụ cấp cố định
    # ---------------------------------------------------------------
    
    # Chi tiết thử việc (Điều 2 - Dành cho HĐ thử việc)
    thoi_gian_thu_viec = fields.Char(string='Thời gian thử việc', default='01 tháng')
    tu_ngay_thu_viec = fields.Date(string='Thử việc từ ngày')
    den_ngay_thu_viec = fields.Date(string='Thử việc đến ngày')

    # Phụ cấp (Điều 6)
    phu_cap_an_trua = fields.Float(string='Phụ cấp ăn trưa', default=800000)
    phu_cap_dien_thoai = fields.Float(string='Phụ cấp điện thoại', default=500000)
    phu_cap_xang_xe = fields.Float(string='Phụ cấp xăng xe', default=500000)

    ghi_chu = fields.Text(string='Ghi chú')

    # Hợp đồng đã ký
    file_hop_dong_da_ky = fields.Binary(string='Hợp đồng đã ký', attachment=True)
    file_hop_dong_name = fields.Char(string='Tên file hợp đồng')
    contract_preview = fields.Html(string='Xem trước mẫu hợp đồng', compute='_compute_contract_preview')



    @api.depends('loai_hop_dong', 'nhan_vien_id', 'ngay_ky', 'dia_diem_ky', 'cong_viec_phai_lam', 'luong_co_ban', 'phu_cap_an_trua', 'phu_cap_dien_thoai', 'phu_cap_xang_xe', 'thoi_gio_lam_viec', 'thoi_gian_nghi')
    def _compute_contract_preview(self):
        for rec in self:
            nv_name = rec.nhan_vien_id.name if rec.nhan_vien_id else "..........................."
            prefix_b = "Ông/Bà"
            if rec.nhan_vien_id:
                prefix_b = "Bà" if rec.gioi_tinh_nv == "nu" else "Ông"
                
            ngay_sinh = rec.ngay_sinh_nv.strftime("%d/%m/%Y") if rec.ngay_sinh_nv else "..."
            
            def get_sel_label(record, field_name):
                val = getattr(record, field_name)
                if not val: return "..."
                field = record._fields[field_name]
                try:
                    sel = field.selection
                    if callable(sel): sel = sel(record)
                    return dict(sel).get(val, str(val))
                except Exception:
                    return str(val)
                    
            gioi_tinh = get_sel_label(rec, "gioi_tinh_nv")
            so_cccd = rec.so_cmnd_nv or "..."
            ngay_cap = rec.ngay_cap_cmnd_nv.strftime("%d/%m/%Y") if rec.ngay_cap_cmnd_nv else "..."
            noi_cap = rec.noi_cap_cmnd_nv or "..."
            hktt = rec.ho_khau_thuong_tru_nv or "................................................................................"
            cho_o = rec.dia_chi_lien_he_nv or "................................................................................"
            chuc_vu = rec.chuc_vu_id.name if rec.chuc_vu_id else "Nhân viên"
            
            if rec.ngay_ky:
                d_str = f"ngày {rec.ngay_ky.day:02d} tháng {rec.ngay_ky.month:02d} năm {rec.ngay_ky.year}"
            else:
                d_str = "ngày ... tháng ... năm 20..."
            dia_diem_ky_text = f"{rec.dia_diem_ky or 'Thành phố Hồ Chí Minh'}, {d_str}"
            
            cong_viec = ""
            if rec.cong_viec_phai_lam:
                for line in rec.cong_viec_phai_lam.split('\n'):
                    if line.strip():
                        cong_viec += f"<p style='margin: 2px 0 2px 30px; text-indent: 0; font-family: \"Times New Roman\", serif; font-size: 14px;'>- {line.strip()}</p>"
            else:
                cong_viec = "<p style='margin: 2px 0 2px 30px; text-indent: 0; font-family: \"Times New Roman\", serif; font-size: 14px;'>- Làm bánh, nhặt bánh, trộn bánh, ...</p>"

            if rec.loai_hop_dong == 'thu_viec':
                chi_tiet_loai_hd = f"""
                    <p style="margin: 4px 0; font-family: 'Times New Roman', serif; font-size: 14px;">Loại hợp đồng: <b>Hợp đồng lao động thử việc</b></p>
                    <p style="margin: 4px 0; font-family: 'Times New Roman', serif; font-size: 14px;">Trong đó có thỏa thuận nội dung thử việc như sau:</p>
                    <p style="margin: 2px 0 2px 20px; font-family: 'Times New Roman', serif; font-size: 14px;">- Thời gian thử việc: {rec.thoi_gian_thu_viec or "01 tháng"}</p>
                    <p style="margin: 2px 0 2px 20px; font-family: 'Times New Roman', serif; font-size: 14px;">- Từ ngày {rec.tu_ngay_thu_viec.strftime("%d/%m/%Y") if rec.tu_ngay_thu_viec else "..."} đến hết {rec.den_ngay_thu_viec.strftime("%d/%m/%Y") if rec.den_ngay_thu_viec else "..."}</p>
                    <p style="margin: 2px 0 2px 20px; font-family: 'Times New Roman', serif; font-size: 14px;">+ Trường hợp thử việc không đạt yêu cầu thì người sử dụng lao động có quyền chấm dứt hợp đồng lao động này.</p>
                    <p style="margin: 2px 0 2px 20px; font-family: 'Times New Roman', serif; font-size: 14px;">+ Trường hợp thử việc đạt yêu cầu thì hai bên tiếp tục thực hiện hợp đồng lao động chính thức.</p>
                """
            else:
                loai_hd_label = get_sel_label(rec, "loai_hop_dong")
                chi_tiet_loai_hd = f"""
                    <p style="margin: 4px 0; font-family: 'Times New Roman', serif; font-size: 14px;">- Loại hợp đồng lao động: <b>{loai_hd_label}</b></p>
                    <p style="margin: 4px 0; font-family: 'Times New Roman', serif; font-size: 14px;">- Bắt đầu từ ngày {rec.ngay_bat_dau.strftime("%d/%m/%Y") if rec.ngay_bat_dau else "..."}</p>
                    <p style="margin: 4px 0; font-family: 'Times New Roman', serif; font-size: 14px;">- Kết thúc vào ngày {rec.ngay_ket_thuc.strftime("%d/%m/%Y") if rec.ngay_ket_thuc else "..."}</p>
                """
            
            luong_cb_val = rec.luong_co_ban or 0.0
            phu_cap_at = rec.phu_cap_an_trua or 0.0
            phu_cap_dt = rec.phu_cap_dien_thoai or 0.0
            phu_cap_xx = rec.phu_cap_xang_xe or 0.0
            
            try_view_probation_wage = f"<p style='margin: 2px 0 2px 15px; font-family: \"Times New Roman\", serif;'>Trong đó: Thời gian thử việc được hưởng mức lương = 85% x Mức lương chính</p>" if rec.loai_hop_dong == 'thu_viec' else ""

            html_content = f"""
            <div style="font-family: 'Times New Roman', Times, serif; font-size: 14px; line-height: 1.5; color: #000; padding: 40px; border: 1px solid #ccc; background-color: #fff; max-width: 800px; margin: 0 auto; box-shadow: 0 4px 8px rgba(0,0,0,0.1);">
                <div style="text-align: center; font-weight: bold; margin-bottom: 5px; font-size: 15px; font-family: 'Times New Roman', serif;">
                    CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
                </div>
                <div style="text-align: center; font-weight: bold; margin-bottom: 5px; font-size: 16px; font-family: 'Times New Roman', serif;">
                    Độc lập - Tự do - Hạnh phúc
                </div>
                <div style="text-align: center; margin-bottom: 15px; font-family: 'Times New Roman', serif;">
                    -----------------------
                </div>
                <div style="text-align: right; font-style: italic; margin-bottom: 25px; font-family: 'Times New Roman', serif;">
                    {dia_diem_ky_text}
                </div>
                
                <div style="margin-left: 40px; font-style: italic; font-size: 13px; margin-bottom: 25px; line-height: 1.4; font-family: 'Times New Roman', serif;">
                    <p style="margin: 2px 0;">- Căn cứ vào Bộ Luật Lao Động số 45/2019/QH14 ngày 20 tháng 11 năm 2019;</p>
                    <p style="margin: 2px 0;">- Căn cứ vào Nghị định 145/2020/NĐ-CP ngày 14 tháng 12 năm 2020 quy định chi tiết và hướng dẫn thi hành một số điều của Bộ luật lao động;</p>
                    <p style="margin: 2px 0;">- Căn cứ vào Thông tư 10/2020/TT-BLĐTBXH ngày 12 tháng 11 năm 2020 hướng dẫn một số nội dung tại Bộ luật Lao động.</p>
                </div>
                
                <div style="text-align: center; font-weight: bold; font-size: 18px; margin-bottom: 5px; font-family: 'Times New Roman', serif;">
                    HỢP ĐỒNG LAO ĐỘNG
                </div>
                <div style="text-align: center; font-style: italic; margin-bottom: 25px; font-family: 'Times New Roman', serif;">
                    (Số: ....................)
                </div>
                
                <p style="margin-bottom: 15px; font-family: 'Times New Roman', serif;">Chúng tôi gồm có:</p>
                
                <p style="font-weight: bold; margin-bottom: 5px; font-family: 'Times New Roman', serif;">Bên người sử dụng lao động: CÔNG TY TNHH THIÊN THỜI</p>
                <div style="margin-left: 15px; margin-bottom: 15px; font-family: 'Times New Roman', serif;">
                    <p style="margin: 2px 0;">Mã số thuế: {rec.mst_a or "..."}</p>
                    <p style="margin: 2px 0;">Địa chỉ: {rec.dia_chi_a or "..."}</p>
                    <p style="margin: 2px 0;">Đại diện bởi: {rec.nguoi_dai_dien_a or "..................."} &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; Chức vụ: {rec.chuc_vu_a or "..........."}</p>
                    <p style="margin: 2px 0; font-style: italic;">(Sau đây gọi tắt là: “NSDLĐ” hoặc “Công ty”)</p>
                </div>
                
                <p style="text-align: center; font-style: italic; margin: 10px 0; font-family: 'Times New Roman', serif;">và</p>
                
                <p style="font-weight: bold; margin-bottom: 5px; font-family: 'Times New Roman', serif;">Bên người lao động: {prefix_b} {nv_name}</p>
                <div style="margin-left: 15px; margin-bottom: 15px; font-family: 'Times New Roman', serif;">
                    <table style="width: 100%; border: none; margin-bottom: 5px; font-family: 'Times New Roman', serif;">
                        <tr style="border: none;">
                            <td style="width: 50%; border: none; padding: 2px 0; font-family: 'Times New Roman', serif;">Sinh ngày: {ngay_sinh}</td>
                            <td style="width: 50%; border: none; padding: 2px 0; font-family: 'Times New Roman', serif;">Giới tính: {gioi_tinh}</td>
                        </tr>
                        <tr style="border: none;">
                            <td style="width: 50%; border: none; padding: 2px 0; font-family: 'Times New Roman', serif;">Số CCCD: {so_cccd}</td>
                            <td style="width: 50%; border: none; padding: 2px 0; font-family: 'Times New Roman', serif;">Cấp ngày: {ngay_cap}</td>
                        </tr>
                    </table>
                    <p style="margin: 2px 0;">Nơi cấp: {noi_cap}</p>
                    <p style="margin: 2px 0;">Nơi đăng ký HKTT: {hktt}</p>
                    <p style="margin: 2px 0;">Chỗ ở hiện nay: {cho_o}</p>
                    <p style="margin: 2px 0; font-style: italic;">(Sau đây gọi tắt là: “NLĐ”)</p>
                </div>
                
                <p style="margin: 15px 0; font-style: italic; font-family: 'Times New Roman', serif;">Thỏa thuận ký kết hợp đồng lao động và cam kết làm đúng những điều khoản sau đây:</p>
                
                <p style="font-weight: bold; margin-top: 15px; margin-bottom: 5px; font-family: 'Times New Roman', serif;">Điều 1. Công việc và địa điểm làm việc:</p>
                <div style="margin-left: 15px; margin-bottom: 15px; font-family: 'Times New Roman', serif;">
                    <p style="margin: 2px 0;">1) Công việc:</p>
                    <p style="margin: 2px 0 15px 15px;">- Vị trí: <b>{chuc_vu}</b></p>
                    <p style="margin: 2px 0 5px 15px;">- Công việc phải làm:</p>
                    {cong_viec}
                    <p style="margin: 10px 0 2px 0;">2) Địa điểm làm việc của người lao động:</p>
                    <p style="margin: 2px 0 2px 15px;">Tại {rec.dia_diem_lam_viec or "Trụ sở công ty"}.</p>
                </div>
                
                <p style="font-weight: bold; margin-top: 15px; margin-bottom: 5px; font-family: 'Times New Roman', serif;">Điều 2: Loại hợp đồng và thời hạn của hợp đồng:</p>
                <div style="margin-left: 15px; margin-bottom: 15px; font-family: 'Times New Roman', serif;">
                    {chi_tiet_loai_hd}
                </div>
                
                <p style="font-weight: bold; margin-top: 15px; margin-bottom: 5px; font-family: 'Times New Roman', serif;">Điều 3: Thời giờ làm việc, thời giờ nghỉ ngơi:</p>
                <div style="margin-left: 15px; margin-bottom: 15px; font-family: 'Times New Roman', serif;">
                    <p style="margin: 2px 0;">1) Thời giờ làm việc: {rec.thoi_gio_lam_viec or "..."}</p>
                    <p style="margin: 2px 0;">2) Thời gian nghỉ: {rec.thoi_gian_nghi or "..."}</p>
                </div>
                
                <p style="font-weight: bold; margin-top: 15px; margin-bottom: 5px; font-family: 'Times New Roman', serif;">Điều 4: Quyền lợi và nghĩa vụ của người lao động</p>
                <div style="margin-left: 15px; margin-bottom: 15px; font-family: 'Times New Roman', serif;">
                    <p style="font-weight: bold; margin: 4px 0;">I. Quyền lợi:</p>
                    <p style="margin: 2px 0;">1. Mức lương:</p>
                    <p style="margin: 2px 0 2px 15px;">- Mức lương chính: <b>{luong_cb_val:,.0f} đồng/tháng</b>. Tính trên 26 ngày công.</p>
                    {try_view_probation_wage}
                    <p style="margin: 4px 0 2px 0;">2. Thời hạn trả lương: Được trả lương vào cuối tháng bằng tiền mặt hoặc chuyển khoản.</p>
                    <p style="margin: 2px 0;">3. Hình thức trả lương: Trả qua tài khoản cá nhân hoặc tiền mặt.</p>
                </div>
                
                <p style="font-weight: bold; margin-top: 15px; margin-bottom: 5px; font-family: 'Times New Roman', serif;">Điều 5: Nghĩa vụ và quyền hạn của người sử dụng lao động</p>
                <p style="margin-left: 15px; margin-bottom: 15px; font-family: 'Times New Roman', serif;">Tuân thủ các quy định của Bộ luật Lao động và quy chế công ty.</p>
                
                <p style="font-weight: bold; margin-top: 15px; margin-bottom: 5px; font-family: 'Times New Roman', serif;">Điều 6: Điều khoản thi hành</p>
                <div style="margin-left: 15px; margin-bottom: 25px; font-family: 'Times New Roman', serif;">
                    <p style="margin: 2px 0;">Hợp đồng này làm thành 02 bản có giá trị ngang nhau, mỗi bên giữ một bản.</p>
                    <p style="margin: 2px 0;">Hợp đồng có hiệu lực từ ngày {rec.ngay_bat_dau.strftime("%d/%m/%Y") if rec.ngay_bat_dau else "..."}.</p>
                </div>
                
                <table style="width: 100%; border: none; margin-top: 30px; font-family: 'Times New Roman', serif;">
                    <tr style="border: none;">
                        <td style="width: 50%; border: none; text-align: center; vertical-align: top; font-family: 'Times New Roman', serif;">
                            <p style="font-weight: bold; margin: 0;">Người lao động</p>
                            <p style="margin: 2px 0; font-size: 12px; font-style: italic;">(Ký tên)</p>
                            <p style="margin: 2px 0; font-size: 12px; font-style: italic;">Ghi rõ Họ và Tên</p>
                        </td>
                        <td style="width: 50%; border: none; text-align: center; vertical-align: top; font-family: 'Times New Roman', serif;">
                            <p style="font-weight: bold; margin: 0;">NGƯỜI SỬ DỤNG LAO ĐỘNG</p>
                            <p style="margin: 2px 0; font-size: 12px; font-style: italic;">(Ký tên, đóng dấu)</p>
                            <p style="margin: 2px 0; font-size: 12px; font-style: italic;">Ghi rõ Họ và tên</p>
                        </td>
                    </tr>
                </table>
            </div>
            """
            rec.contract_preview = html_content


    
    @api.model
    def default_get(self, fields):
        res = super(HopDongLaoDong, self).default_get(fields)
        if 'name' in fields and res.get('name', '/') == '/':
            res['name'] = self.env['ir.sequence'].next_by_code('nhan_su.hop_dong') or '/'
        
        if 'contract_preview' in fields:
            loai_hd = res.get('loai_hop_dong') or self._context.get('default_loai_hop_dong') or '36_thang'
            tmp_rec = self.new({
                'loai_hop_dong': loai_hd,
                'name': res.get('name', '....'),
            })
            try:
                tmp_rec._compute_working_hours()
            except Exception:
                pass
            try:
                tmp_rec._compute_luong_tu_cau_hinh()
            except Exception:
                pass
            try:
                tmp_rec._compute_contract_preview()
            except Exception:
                pass
            res['contract_preview'] = tmp_rec.contract_preview
        return res

    def action_upload_signed_contract(self):
        """Mở wizard upload"""
        self.ensure_one()
        return {
            'name': 'Upload Hợp đồng đã ký',
            'type': 'ir.actions.act_window',
            'res_model': 'nhan_su.hop_dong.upload',
            'view_mode': 'form',
            'target': 'new',
            'context': {'default_hop_dong_id': self.id},
        }

    def action_activate(self):
        self.write({'state': 'active'})

    def action_cancel(self):
        self.write({'state': 'cancel'})

    def action_set_draft(self):
        self.write({'state': 'draft'})

    @api.model
    def create(self, vals):
        if vals.get('name', '/') == '/':
            existing_contracts = self.search([('name', '=like', 'HD%')])
            numbers = []
            for c in existing_contracts:
                if c.name:
                    try:
                        num_str = ''.join(filter(str.isdigit, c.name))
                        if num_str:
                            numbers.append(int(num_str))
                    except Exception:
                        pass
            next_number = max(numbers) + 1 if numbers else 1
            vals['name'] = f'HD{next_number:04d}'
        return super(HopDongLaoDong, self).create(vals)
    
    def unlink(self):
        return super(HopDongLaoDong, self).unlink()
    
    @api.model
    def cron_update_contract_state(self):
        """Tự động cảnh báo sắp hết hạn hợp đồng"""
        today = fields.Date.context_today(self)
        
        # 1. Tạo cảnh báo (Activity) cho các hợp đồng sắp hết hạn
        warning_official = today + relativedelta(days=30)
        warning_probation = today + relativedelta(days=7)

        contracts_to_warn = self.search([
            ('ngay_ket_thuc', '!=', False),
            '|',
            '&', ('loai_hop_dong', '!=', 'thu_viec'), ('ngay_ket_thuc', '<=', warning_official),
            '&', ('loai_hop_dong', '=', 'thu_viec'), ('ngay_ket_thuc', '<=', warning_probation),
            ('ngay_ket_thuc', '>=', today) # Chỉ cảnh báo những cái chưa hết hạn hẳn
        ])

        for contract in contracts_to_warn:
            if contract.state == 'active':
                contract.write({'state': 'expiring'})
            existing_activity = self.env['mail.activity'].search([
                ('res_id', '=', contract.id),
                ('res_model_id', '=', self.env['ir.model']._get_id('nhan_su.hop_dong')),
                ('summary', 'like', 'Sắp hết hạn hợp đồng')
            ])
            if not existing_activity:
                nv_name_activity = contract.nhan_vien_id.name if contract.nhan_vien_id else "Chưa xác định"
                contract.activity_schedule(
                    'mail.mail_activity_data_todo',
                    summary=f'Sắp hết hạn hợp đồng: {nv_name_activity}',
                    note=f'Hợp đồng {contract.name} sẽ hết hạn vào ngày {contract.ngay_ket_thuc}. Vui lòng chuẩn bị gia hạn.',
                    date_deadline=contract.ngay_ket_thuc,
                    user_id=self.env.user.id
                )
        
        # 2. Tự động chuyển trạng thái sang 'expired' nếu đã quá ngày kết thúc
        expired_contracts = self.search([
            ('state', 'in', ['active', 'expiring']),
            ('ngay_ket_thuc', '<', today)
        ])
        if expired_contracts:
            expired_contracts.write({'state': 'expired'})
            
        return "Cron executed successfully."
    
    def get_active_contract(self, date_from):
        """Lấy hợp đồng đang hoạt động vào ngày chỉ định"""
        return self.search([
            ('nhan_vien_id', '=', self.nhan_vien_id.id),
            ('ngay_bat_dau', '<=', date_from),
            ('state', '=', 'active'),
            '|',
            ('ngay_ket_thuc', '=', False),
            ('ngay_ket_thuc', '>=', date_from),
        ], limit=1)

    def action_export_docx(self):
        self.ensure_one()
        if not Document:
            return {
                'type': 'ir.actions.client', 'tag': 'display_notification',
                'params': {'title': _('Lỗi'), 'message': _('Thư viện python-docx chưa được cài đặt.'), 'type': 'danger'}
            }

        document = Document()
        
        # --- SETUP PAGE & STYLE ---
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)
        
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_after = Pt(6)
        
        def get_sel_label(record, field_name):
            val = getattr(record, field_name)
            if not val: return "..."
            field = record._fields[field_name]
            try:
                sel = field.selection
                if callable(sel): sel = sel(record)
                return dict(sel).get(val, str(val))
            except Exception:
                return str(val)
        
        # Quốc hiệu
        p_qh = document.add_paragraph()
        p_qh.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p_qh.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
        run1.bold = True; run1.font.size = Pt(13)
        p_qh2 = document.add_paragraph()
        p_qh2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p_qh2.add_run('Độc lập - Tự do - Hạnh phúc')
        run2.bold = True; run2.font.size = Pt(14)
        p_qh3 = document.add_paragraph()
        p_qh3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_qh3.add_run('-----------------------')

        # Ngày tháng (Lệch phải) - Format mới
        p_date = document.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if self.ngay_ky:
            d_str = f'ngày {self.ngay_ky.day} tháng {self.ngay_ky.month} năm {self.ngay_ky.year}'
        else:
            d_str = 'ngày ... tháng ... năm 2026'
        p_date.add_run(f'{self.dia_diem_ky or "Thành phố Hồ Chí Minh"}, {d_str}')

        # Căn cứ (Thụt lề, In nghiêng, Thẳng hàng)
        document.add_paragraph()  # Thêm dòng trống tạo giãn cách đẹp mắt
        for text in [
            'Căn cứ vào Bộ Luật Lao Động số 45/2019/QH14 ngày 20 tháng 11 năm 2019;',
            'Căn cứ vào Nghị định 145/2020/NĐ-CP ngày 14 tháng 12 năm 2020 quy định chi tiết và hướng dẫn thi hành một số điều của Bộ luật lao động;',
            'Căn cứ vào Thông tư 10/2020/TT-BLĐTBXH ngày 12 tháng 11 năm 2020 hướng dẫn một số nội dung tại Bộ luật Lao động.'
        ]:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f'- {text}')
            run.italic = True

        # Tiêu đề
        p_title = document.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run('\nHỢP ĐỒNG LAO ĐỘNG')
        run_title.bold = True; run_title.font.size = Pt(16)
        p_so = document.add_paragraph()
        p_so.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_so.add_run('(Số: ....................)')

        document.add_paragraph('\nChúng tôi gồm có:')

        # Bên A
        p_a = document.add_paragraph()
        p_a.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_a = p_a.add_run('Bên người sử dụng lao động: CÔNG TY TNHH THIÊN THỜI')
        run_a.bold = True
        
        p_mst = document.add_paragraph(f'Mã số thuế: {self.mst_a or "..."}')
        p_mst.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_dia_chi = document.add_paragraph(f'Địa chỉ: {self.dia_chi_a or "..."}')
        p_dia_chi.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_rep = document.add_paragraph(f'Đại diện bởi: {self.nguoi_dai_dien_a or "..................."}\tChức vụ: {self.chuc_vu_a or "..........."}')
        p_rep.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_nsdld = document.add_paragraph('(Sau đây gọi tắt là: “NSDLĐ” hoặc “Công ty”)')
        p_nsdld.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_va = document.add_paragraph('\nvà')
        p_va.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Bên B
        p_b = document.add_paragraph()
        p_b.alignment = WD_ALIGN_PARAGRAPH.LEFT
        nv_name = self.nhan_vien_id.name if self.nhan_vien_id else "..........................."
        prefix_b = "Ông/Bà"
        if self.nhan_vien_id:
             prefix_b = "Bà" if self.gioi_tinh_nv == "nu" else "Ông"
        
        run_b = p_b.add_run(f'Bên người lao động: {prefix_b} {nv_name}')
        run_b.bold = True
        
        # Sinh ngày & Giới tính trên cùng 1 dòng
        p_birth = document.add_paragraph(f'Sinh ngày: {self.ngay_sinh_nv.strftime("%d/%m/%Y") if self.ngay_sinh_nv else "..."}\t\t\tGiới tính: {get_sel_label(self, "gioi_tinh_nv")}')
        p_birth.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # CCCD & Ngày cấp trên cùng 1 dòng
        p_cccd = document.add_paragraph(f'Số CCCD: {self.so_cmnd_nv or "..."}\t\t\tCấp ngày: {self.ngay_cap_cmnd_nv.strftime("%d/%m/%Y") if self.ngay_cap_cmnd_nv else "..."}')
        p_cccd.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_noi_cap = document.add_paragraph(f'Nơi cấp: {self.noi_cap_cmnd_nv or "..."}')
        p_noi_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_hktt = document.add_paragraph('Nơi đăng ký HKTT: .......................................................................................................................')
        p_hktt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_cho_o = document.add_paragraph('Chỗ ở hiện nay: .......................................................................................................................')
        p_cho_o.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_nld = document.add_paragraph('(Sau đây gọi tắt là: “NLĐ”)')
        p_nld.alignment = WD_ALIGN_PARAGRAPH.LEFT

        document.add_paragraph('\nThỏa thuận ký kết hợp đồng lao động và cam kết làm đúng những điều khoản sau đây:')

        # Điều 1
        p1 = document.add_paragraph(); p1.add_run('Điều 1. Công việc và địa điểm làm việc:').bold = True
        document.add_paragraph('1) Công việc:').paragraph_format.left_indent = Inches(0.2)
        document.add_paragraph(f'- Vị trí: {self.chuc_vu_id.name or "Nhân viên"}').paragraph_format.left_indent = Inches(0.3)
        document.add_paragraph('- Công việc phải làm:').paragraph_format.left_indent = Inches(0.3)
        for line in (self.cong_viec_phai_lam or "").split('\n'):
            if line.strip():
                document.add_paragraph(line.strip()).paragraph_format.left_indent = Inches(0.4)
        document.add_paragraph(f'2) Địa điểm làm việc của người lao động:').paragraph_format.left_indent = Inches(0.2)
        document.add_paragraph(f'Tại {self.dia_diem_lam_viec or "Trụ sở công ty"}.').paragraph_format.left_indent = Inches(0.3)

        # Điều 2
        p2 = document.add_paragraph(); p2.add_run('Điều 2: Loại hợp đồng và thời hạn của hợp đồng:').bold = True
        if self.loai_hop_dong == 'thu_viec':
            document.add_paragraph(f'Loại hợp đồng: Hợp đồng lao động thử việc')
            document.add_paragraph(f'Trong đó có thỏa thuận nội dung thử việc như sau:')
            document.add_paragraph(f'- Thời gian thử việc: {self.thoi_gian_thu_viec or "01 tháng"}').paragraph_format.left_indent = Inches(0.2)
            d_tu = self.tu_ngay_thu_viec.strftime("%d/%m/%Y") if self.tu_ngay_thu_viec else "..."
            d_den = self.den_ngay_thu_viec.strftime("%d/%m/%Y") if self.den_ngay_thu_viec else "..."
            document.add_paragraph(f'- Từ ngày {d_tu} đến hết {d_den}').paragraph_format.left_indent = Inches(0.2)
            document.add_paragraph('+ Trường hợp thử việc không đạt yêu cầu thì người sử dụng lao động có quyền chấm dứt hợp đồng lao động này.').paragraph_format.left_indent = Inches(0.2)
            document.add_paragraph('+ Trường hợp thử việc đạt yêu cầu thì hai bên tiếp tục thực hiện hợp đồng lao động chính thức.').paragraph_format.left_indent = Inches(0.2)
        else:
            document.add_paragraph(f'- Loại hợp đồng lao động: {get_sel_label(self, "loai_hop_dong")}')
            document.add_paragraph(f'- Bắt đầu từ ngày {self.ngay_bat_dau.strftime("%d/%m/%Y") if self.ngay_bat_dau else "..."}')
            document.add_paragraph(f'- Kết thúc vào ngày {self.ngay_ket_thuc.strftime("%d/%m/%Y") if self.ngay_ket_thuc else "..."}')

        # Điều 3
        p3 = document.add_paragraph(); p3.add_run('Điều 3: Thời giờ làm việc, thời giờ nghỉ ngơi:').bold = True
        document.add_paragraph(f'1) Thời giờ làm việc: {self.thoi_gio_lam_viec or "..."}')
        ca = self.nhan_vien_id.ca_lam_id if self.nhan_vien_id else False
        if ca:
            document.add_paragraph(f'- Trong tuần: {ca.so_ngay_tuan} ngày/tuần: từ {get_sel_label(ca, "tu_thu")} đến {get_sel_label(ca, "den_thu")}')
        document.add_paragraph(f'2) Thời gian nghỉ: {self.thoi_gian_nghi or "..."}')

        # Điều 4
        p4 = document.add_paragraph(); p4.add_run('Điều 4: Quyền lợi và nghĩa vụ của người lao động').bold = True
        document.add_paragraph('I. Quyền lợi:').bold = True
        document.add_paragraph('1. Mức lương:').paragraph_format.left_indent = Inches(0.2)
        document.add_paragraph(f'- Mức lương chính: {self.luong_co_ban:,.0f} đồng/tháng. Tính trên 26 ngày công.').paragraph_format.left_indent = Inches(0.3)
        if self.loai_hop_dong == 'thu_viec':
            document.add_paragraph('Trong đó:').paragraph_format.left_indent = Inches(0.3)
            document.add_paragraph('Thời gian thử việc được hưởng mức lương = 85% x Mức lương chính').paragraph_format.left_indent = Inches(0.4)
        
        document.add_paragraph('2. Thời hạn trả lương: Được trả lương vào cuối tháng bằng tiền mặt hoặc chuyển khoản.').paragraph_format.left_indent = Inches(0.2)
        document.add_paragraph('3. Hình thức trả lương: Trả qua tài khoản cá nhân hoặc tiền mặt.').paragraph_format.left_indent = Inches(0.2)

        # Điều 5
        p5 = document.add_paragraph(); p5.add_run('Điều 5: Nghĩa vụ và quyền hạn của người sử dụng lao động').bold = True
        document.add_paragraph('Tuân thủ các quy định của Bộ luật Lao động và quy chế công ty.')

        # Điều 6
        p6 = document.add_paragraph(); p6.add_run('Điều 6: Điều khoản thi hành').bold = True
        document.add_paragraph(f'Hợp đồng này làm thành 02 bản có giá trị ngang nhau, mỗi bên giữ một bản.')
        document.add_paragraph(f'Hợp đồng có hiệu lực từ ngày {self.ngay_bat_dau.strftime("%d/%m/%Y") if self.ngay_bat_dau else "..."}.')

        # Chữ ký
        document.add_paragraph()
        table = document.add_table(rows=1, cols=2)
        c1, c2 = table.rows[0].cells
        p_c1 = c1.paragraphs[0]; p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_c1 = p_c1.add_run('Người lao động'); run_c1.bold = True
        c1.add_paragraph('(Ký tên)').alignment = WD_ALIGN_PARAGRAPH.CENTER
        c1.add_paragraph('Ghi rõ Họ và Tên').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_c2 = c2.paragraphs[0]; p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_c2 = p_c2.add_run('NGƯỜI SỬ DỤNG LAO ĐỘNG'); run_c2.bold = True
        c2.add_paragraph('(Ký tên, đóng dấu)').alignment = WD_ALIGN_PARAGRAPH.CENTER
        c2.add_paragraph('Ghi rõ Họ và tên').alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save
        fp = io.BytesIO(); document.save(fp); fp.seek(0)
        data = base64.b64encode(fp.read()); fp.close()
        attachment = self.env['ir.attachment'].create({
            'name': f'Hop_Dong_{self.name}.docx', 'datas': data,
            'res_model': 'nhan_su.hop_dong', 'res_id': self.id, 'type': 'binary',
        })
        return {'type': 'ir.actions.act_url', 'url': f'/web/content/{attachment.id}?download=true', 'target': 'new'}
